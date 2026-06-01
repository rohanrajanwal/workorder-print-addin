"""
Build the Print Work Order Add-In — Technical Deep Dive Word doc.
Uses python-docx with manual styling to produce a polished, branded document.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(r"C:\Users\rohandeeprajanwal\ai-projects\workorder-print-addin\docs\Print-Work-Order-Technical-Deep-Dive.docx")

# Brand palette — Geotab-ish blue (matches the print button)
BRAND_BLUE = RGBColor(0x2D, 0x7F, 0xF9)
BRAND_BLUE_DARK = RGBColor(0x1A, 0x6C, 0xE5)
INK = RGBColor(0x1A, 0x1A, 0x1A)
GREY_DARK = RGBColor(0x44, 0x44, 0x44)
GREY_MID = RGBColor(0x66, 0x66, 0x66)
GREY_LIGHT = RGBColor(0xBB, 0xBB, 0xBB)
CODE_BG = "F4F6F8"
HEADER_FILL = "E8F1FE"  # very light blue
DIVIDER = RGBColor(0xE0, 0xE0, 0xE0)

doc = Document()

# ---- Page setup: US Letter, 1" margins ----
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ---- Default styles ----
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

def set_font_for_style(style_name, *, name='Calibri', size=11, bold=False, color=INK,
                      space_before=0, space_after=6):
    s = styles[style_name]
    s.font.name = name
    s.font.size = Pt(size)
    s.font.bold = bold
    s.font.color.rgb = color
    s.paragraph_format.space_before = Pt(space_before)
    s.paragraph_format.space_after = Pt(space_after)
    s.paragraph_format.keep_with_next = True

set_font_for_style('Heading 1', name='Calibri', size=22, bold=True, color=BRAND_BLUE_DARK,
                   space_before=18, space_after=6)
set_font_for_style('Heading 2', name='Calibri', size=16, bold=True, color=INK,
                   space_before=14, space_after=4)
set_font_for_style('Heading 3', name='Calibri', size=13, bold=True, color=GREY_DARK,
                   space_before=10, space_after=2)
set_font_for_style('Title', name='Calibri', size=36, bold=True, color=BRAND_BLUE_DARK,
                   space_before=0, space_after=4)
set_font_for_style('Subtitle', name='Calibri', size=14, bold=False, color=GREY_MID,
                   space_before=0, space_after=24)

# ---- Helpers ----
def add_para(text='', *, bold=False, italic=False, size=11, color=INK,
             align=None, space_before=0, space_after=6, font='Calibri'):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.name = font
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
    return p

def add_runs(parts, *, align=None, space_after=6, space_before=0):
    """parts: list of (text, dict-of-font-props) tuples."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for text, props in parts:
        r = p.add_run(text)
        r.font.name = props.get('font', 'Calibri')
        r.font.size = Pt(props.get('size', 11))
        r.font.bold = props.get('bold', False)
        r.font.italic = props.get('italic', False)
        r.font.color.rgb = props.get('color', INK)
    return p

def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC", size="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        borders.append(b)
    tcPr.append(borders)

def cell_text(cell, text, *, bold=False, color=INK, size=10.5, font='Calibri',
              align=None, italic=False):
    cell.text = ''
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_code_block(code, *, caption=None):
    """Single-cell shaded table that renders as a code block."""
    if caption:
        add_runs([(caption, {'size': 10, 'italic': True, 'color': GREY_MID})],
                 space_after=2)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)
    shade_cell(cell, CODE_BG)
    set_cell_borders(cell, color="DDDDDD", size="4")
    cell.text = ''
    for i, line in enumerate(code.split('\n')):
        p = cell.add_paragraph() if i > 0 else cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(line if line else ' ')
        r.font.name = 'Consolas'
        r.font.size = Pt(9.5)
        r.font.color.rgb = INK
    # trailing space
    add_para('', size=4, space_after=8)

def add_bullet(text, *, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent_level)
    p.paragraph_format.space_after = Pt(3)
    r = p.runs[0] if p.runs else p.add_run()
    p.text = ''  # clear default
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.color.rgb = INK
    return p

def add_bullet_mixed(parts):
    """Bullet with mixed-format parts: list of (text, dict)."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    p.text = ''
    for text, props in parts:
        r = p.add_run(text)
        r.font.name = props.get('font', 'Calibri')
        r.font.size = Pt(props.get('size', 11))
        r.font.bold = props.get('bold', False)
        r.font.italic = props.get('italic', False)
        r.font.color.rgb = props.get('color', INK)
    return p

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:color'), '2D7FF9')
    bot.set(qn('w:space'), '1')
    pBdr.append(bot)
    pPr.append(pBdr)

def add_callout(label, body, color_hex="E8F1FE", label_color=BRAND_BLUE_DARK):
    """Boxed callout — single-cell shaded table with a label + body."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)
    shade_cell(cell, color_hex)
    set_cell_borders(cell, color="C5DCFB", size="6")
    cell.text = ''
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(label)
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10)
    r1.font.bold = True
    r1.font.color.rgb = label_color
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(body)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)
    r2.font.color.rgb = INK
    add_para('', size=4, space_after=4)

def add_two_col_table(rows, *, col1_width=Inches(2.0), col2_width=Inches(4.5),
                     header=None):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = col1_width
    tbl.columns[1].width = col2_width
    if header:
        hr = tbl.add_row()
        for i, h in enumerate(header):
            c = hr.cells[i]
            c.width = col1_width if i == 0 else col2_width
            cell_text(c, h, bold=True, color=BRAND_BLUE_DARK, size=10.5)
            shade_cell(c, HEADER_FILL)
            set_cell_borders(c, color="C5DCFB", size="4")
    for row in rows:
        tr = tbl.add_row()
        for i, val in enumerate(row):
            c = tr.cells[i]
            c.width = col1_width if i == 0 else col2_width
            cell_text(c, val, size=10.5, bold=(i == 0))
            set_cell_borders(c, color="E0E0E0", size="4")
    add_para('', size=4, space_after=8)

def add_three_col_table(rows, *, widths=(Inches(1.6), Inches(2.4), Inches(2.5)), header=None):
    tbl = doc.add_table(rows=0, cols=3)
    tbl.autofit = False
    for i, w in enumerate(widths):
        tbl.columns[i].width = w
    if header:
        hr = tbl.add_row()
        for i, h in enumerate(header):
            c = hr.cells[i]
            c.width = widths[i]
            cell_text(c, h, bold=True, color=BRAND_BLUE_DARK, size=10.5)
            shade_cell(c, HEADER_FILL)
            set_cell_borders(c, color="C5DCFB", size="4")
    for row in rows:
        tr = tbl.add_row()
        for i, val in enumerate(row):
            c = tr.cells[i]
            c.width = widths[i]
            cell_text(c, val, size=10.5, bold=(i == 0))
            set_cell_borders(c, color="E0E0E0", size="4")
    add_para('', size=4, space_after=8)

def heading(text, level=1):
    p = doc.add_heading('', level=level)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    if level == 1:
        r.font.size = Pt(22)
        r.font.color.rgb = BRAND_BLUE_DARK
    elif level == 2:
        r.font.size = Pt(15)
        r.font.color.rgb = INK
    elif level == 3:
        r.font.size = Pt(12.5)
        r.font.color.rgb = GREY_DARK
    r.font.bold = True
    r.font.name = 'Calibri'
    return p

# ============================================================
# COVER PAGE
# ============================================================
# spacer to push title down
for _ in range(4):
    add_para('', space_after=0)

add_para("PRINT WORK ORDER", size=11, bold=True, color=BRAND_BLUE,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(8)
r = p.add_run("Technical Deep Dive")
r.font.name = 'Calibri'
r.font.size = Pt(40)
r.font.bold = True
r.font.color.rgb = INK

add_para("How the MyGeotab add-in is built — explained from scratch",
         size=14, color=GREY_MID, space_after=24)

# Coloured divider
add_divider()

# Metadata table on cover
meta = doc.add_table(rows=4, cols=2)
meta.autofit = False
meta.columns[0].width = Inches(1.5)
meta.columns[1].width = Inches(5.0)
for i, (k, v) in enumerate([
    ("Project", "Print Work Order — MyGeotab Add-In"),
    ("Repository", "rohanrajanwal/workorder-print-addin"),
    ("Version", "1.1"),
    ("Document Date", "May 2026"),
]):
    cell_text(meta.rows[i].cells[0], k, bold=True, color=GREY_MID, size=10)
    cell_text(meta.rows[i].cells[1], v, size=10.5)

# Page break to start content
doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
heading("Contents", level=1)

toc_items = [
    "1.  The Big Picture: What Problem Does This Solve?",
    "2.  What Is a MyGeotab Add-In?",
    "3.  The Configuration JSON, Decoded",
    "4.  The Two-Part Install",
    "5.  The JavaScript File: The IIFE Wrapper",
    "6.  The MyGeotab Button-Click Contract",
    "7.  Fetching the Data",
    "8.  The Data Quirks",
    "9.  Building the HTML",
    "10.  CSS and Print-Specific Styling",
    "11.  Triggering the Print",
    "12.  Keeping the Button Styled (MutationObserver)",
    "13.  Shop Info: localStorage + API + Prompt Fallback",
    "14.  The End-to-End Flow",
    "15.  The Dev Harness",
    "16.  Tech Stack Summary",
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.1)
    r = p.add_run(item)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.color.rgb = INK

doc.add_page_break()

# ============================================================
# SECTION 1
# ============================================================
heading("1. The Big Picture: What Problem Does This Solve?", level=1)

add_para("MyGeotab has a Work Order Details page. Out of the box, there is no print button. "
         "If a mechanic wants a paper copy of a work order to take to the shop floor, they have "
         "to manually copy data or screenshot the page — ugly and error-prone.")

add_para("This add-in puts a Print button in the toolbar that generates a clean, formatted "
         "Work Order PDF on demand. The whole thing runs inside the user’s browser — no server, "
         "no database, no backend service. Just JavaScript that ships alongside MyGeotab.")

add_callout("Key idea",
            "An add-in is just JavaScript that MyGeotab loads into specific pages when an "
            "admin tells it to. Once installed, the Print button appears automatically on every "
            "Work Order Details page — for every user in that database.")

# ============================================================
# SECTION 2
# ============================================================
heading("2. What Is a MyGeotab Add-In?", level=1)

add_para("Think of MyGeotab as a website. The Add-In system is MyGeotab’s official way to "
         "inject your own JavaScript into specific pages of that website. It is like a Chrome "
         "extension, but managed by Geotab’s admin panel instead of the Chrome Web Store.")

add_para("When you install an Add-In, you give MyGeotab two things:")
add_bullet_mixed([
    ("A configuration JSON", {'bold': True}),
    (" — tells MyGeotab where (which page) and how (button? menu item? full panel?) "
     "to inject your code.", {})
])
add_bullet_mixed([
    ("One or more JavaScript files", {'bold': True}),
    (" — your actual code.", {})
])

add_para("MyGeotab’s server stores these. When a user opens the matching page, MyGeotab’s "
         "frontend:", space_after=4)
add_bullet("Reads the config")
add_bullet("Sees “there’s an add-in that wants to add a button to maintenanceWorkOrderDetails”")
add_bullet("Creates the button in the toolbar")
add_bullet("When clicked, loads + runs the JavaScript file")

add_para("That’s it. No magic.", italic=True, color=GREY_MID, space_before=6)

# ============================================================
# SECTION 3
# ============================================================
heading("3. The Configuration JSON, Decoded", level=1)

add_code_block('''{
  "name": "Print Work Order",
  "supportEmail": "rohandeep.rajanwal@geotab.com",
  "version": "1.0",
  "items": [{
    "page": "maintenanceWorkOrderDetails",
    "click": "printWorkOrder.js",
    "buttonName": { "en": "Print" }
  }],
  "isSigned": false
}''')

heading("Field-by-field", level=3)
add_two_col_table([
    ("name", "Display name in System Settings → Add-Ins list."),
    ("supportEmail", "Shown to users if they need to report issues."),
    ("version", "Your version number. Cosmetic."),
    ("items", "Array of things this add-in adds. Multiple allowed; we have one button."),
    ("items[].page", "MyGeotab page name where this activates. maintenanceWorkOrderDetails "
                     "is the internal name for the WO Details page. Get it wrong and the button "
                     "never appears."),
    ("items[].click", "Relative filename of the JS file to run when the button is clicked. "
                      "Relative — meaning MyGeotab looks for printWorkOrder.js among the files "
                      "you uploaded with this add-in."),
    ("items[].buttonName", "The label on the button. Object keyed by locale (en, fr, es...) "
                           "for translation support."),
    ("isSigned", "false means this add-in isn’t cryptographically signed by Geotab. "
                 "A MyGeotab admin has to flip “Allow unverified Add-Ins: Yes” for this to run."),
], col1_width=Inches(1.7), col2_width=Inches(4.8))

# ============================================================
# SECTION 4
# ============================================================
heading("4. The Two-Part Install", level=1)

add_para("The MyGeotab UI has a modal with two tabs:")
add_bullet_mixed([("Configuration tab", {'bold': True}), (" — paste the JSON above.", {})])
add_bullet_mixed([("Files tab", {'bold': True}), (" — upload printWorkOrder.js.", {})])

add_para("Both pieces must be uploaded together. The config references files by relative name; "
         "the Files tab is where those files actually live. Internally, MyGeotab stores them as "
         "a bundle.")

add_callout("Why this matters",
            "An earlier attempt tried to install the add-in via the MyGeotab API, passing a "
            "config that referenced a JavaScript URL hosted on GitHub Pages. That failed — "
            "MyGeotab treats external URLs as untrusted. The UI upload route avoids the URL "
            "question entirely: files are uploaded directly to MyGeotab and served from there.")

# ============================================================
# SECTION 5
# ============================================================
heading("5. The JavaScript File: The IIFE Wrapper", level=1)

add_para("The whole printWorkOrder.js is wrapped like this:")
add_code_block('''(function () {
  'use strict';
  // ... everything ...
})();''')

add_para("This is an IIFE (Immediately Invoked Function Expression) — a JavaScript pattern "
         "for scope isolation.")

add_para("Your code runs in the same JavaScript context as MyGeotab itself. If you declared "
         "let api = ... at the top level, you could clash with a variable MyGeotab already "
         "uses. The IIFE wraps everything in a private function scope, so your variables can’t "
         "leak out and pollute the global namespace.")

add_para("The only things this script intentionally exposes are:")
add_bullet_mixed([("window.geotab.customButtons.printWorkOrder",
                   {'font': 'Consolas', 'size': 10.5}),
                  (" — the function MyGeotab calls when the button is clicked", {})])
add_bullet_mixed([("window._woPrint", {'font': 'Consolas', 'size': 10.5}),
                  (" — a dev harness for testing/debugging in the browser console", {})])
add_para("Everything else is private.", italic=True, color=GREY_MID, space_before=4)

# ============================================================
# SECTION 6
# ============================================================
heading("6. The MyGeotab Button-Click Contract", level=1)

add_para("MyGeotab’s add-in system has a specific convention for button add-ins. You assign a "
         "function to window.geotab.customButtons.<yourFunctionName>, where <yourFunctionName> "
         "matches the click field in your config (minus the .js extension).")

add_para("When the button is clicked, MyGeotab calls your function with three arguments:")
add_code_block('''window.geotab.customButtons.printWorkOrder = async function (event, nativeApi, state) {
  // event     — the DOM click event
  // nativeApi — a pre-authenticated MyGeotab API client
  // state     — info about the current page (incl. the WO id you're viewing)
};''')

add_callout("The big gift here",
            "MyGeotab hands you nativeApi — an API client that’s already logged in as the "
            "current user. You don’t deal with cookies, tokens, OAuth, or sessions. Just call "
            "nativeApi.call('Get', { typeName: 'MaintenanceWorkOrder', ... }, success, failure).")

add_para("The catch: nativeApi is callback-based, not Promise-based. Modern JavaScript loves "
         "async/await, so we wrap it:", space_before=4)
add_code_block('''function makePromiseApi(nativeApi) {
  return {
    call: function (method, params) {
      return new Promise(function (resolve, reject) {
        nativeApi.call(method, params, resolve, reject);
      });
    },
    multiCall: function (calls) {
      return new Promise(function (resolve, reject) {
        nativeApi.multiCall(calls, resolve, reject);
      });
    }
  };
}''')

add_para("Now we can write const data = await api.call('Get', ...) instead of nested callbacks.")

# ============================================================
# SECTION 7
# ============================================================
heading("7. Fetching the Data", level=1)

add_code_block('''const [workOrders, jobs] = await api.multiCall([
  ['Get', { typeName: 'MaintenanceWorkOrder', search: { id: woId } }],
  ['Get', { typeName: 'MaintenanceWorkOrderJob', search: { workOrderId: woId } }]
]);''')

add_bullet_mixed([("multiCall", {'bold': True, 'font': 'Consolas', 'size': 10.5}),
                  (" batches multiple API calls into a single network request. Faster than two "
                   "sequential calls.", {})])
add_bullet_mixed([("Get", {'bold': True, 'font': 'Consolas', 'size': 10.5}),
                  (" is the universal MyGeotab read method — works for any entity type.", {})])
add_bullet_mixed([("typeName", {'bold': True, 'font': 'Consolas', 'size': 10.5}),
                  (" specifies which entity. MaintenanceWorkOrder is the WO itself; "
                   "MaintenanceWorkOrderJob is each line item (parts, labor) inside it.", {})])
add_bullet_mixed([("search", {'bold': True, 'font': 'Consolas', 'size': 10.5}),
                  (" is the filter object. The shape depends on the entity.", {})])

add_callout("⚠ The big trap we hit",
            "For MaintenanceWorkOrderJob, the obvious filter search.workOrder.id is silently "
            "ignored by the API. The API doesn’t error — it just returns every job in the "
            "entire database (258 jobs in our test). The only filter that works is the flat-key "
            "form: search.workOrderId. The Geotab SDK docs don’t mention this anywhere; we "
            "discovered it by reverse-engineering MyGeotab’s own UI requests.",
            color_hex="FFF4E5", label_color=RGBColor(0x9A, 0x5B, 0x00))

add_para("Then we fetch the Device separately:", space_before=4)
add_code_block('''const device = (await api.call('Get', {
  typeName: 'Device',
  search: { id: wo.device.id }
}))[0] || {};''')

add_para("Why? The Work Order only stores a reference to the device ({ id: \"b117\" }), not "
         "the full record. To get the VIN, license plate, etc., we fetch the Device entity "
         "using that id.")

# ============================================================
# SECTION 8
# ============================================================
heading("8. The Data Quirks", level=1)

add_para("Once we had the data, three more surprises:")

heading("Quirk 1 — Odometer is hex-encoded", level=3)
add_para("wo.odometerReading looks like \"0000000008246960\". Looks decimal, right? It’s not. "
         "It’s hex. 0x8246960 = 136,605,536. That’s meters. Divide by 1000 → 136,605 km.")
add_code_block('''function fmtOdometer(hexMeters) {
  const meters = parseInt(hexMeters, 16);  // base 16 = hex
  return Math.round(meters / 1000).toLocaleString() + ' km';
}''')
add_para("There is no odometerReadingInMeters field with a decimal value. The hex string is "
         "all you get.", italic=True, color=GREY_MID)

heading("Quirk 2 — Engine hours come from the WO, not the jobs", level=3)
add_para("The WO entity has engineHoursReadingInHours: 5698 — a clean decimal number. The "
         "older version of the script was incorrectly reading this from jobs[0]. Combined with "
         "the broken filter, that meant we were reading engine hours off some random job from a "
         "totally different work order — hence the wildly wrong “1208 hrs” output.")

heading("Quirk 3 — openedByUser is polymorphic", level=3)
add_para("Sometimes it’s a User object {id, name, ...}. Sometimes it’s a raw string ID like "
         "\"b123\". Sometimes it’s the literal sentinel \"NoUserId\" (for WOs created by demo "
         "seeds or automated rules). The script handles all three:")
add_code_block('''const openedByName = wo.openedByUser && typeof wo.openedByUser === 'object'
  ? wo.openedByUser.name
  : null;
const openedBy = openedByName
  ? escHtml(openedByName)
  : (wo.openedByUser && wo.openedByUser !== 'NoUserId'
      ? escHtml(wo.openedByUser)
      : missing('not recorded'));''')

# ============================================================
# SECTION 9
# ============================================================
heading("9. Building the HTML", level=1)

add_para("buildRepairOrderHTML() constructs an HTML string from the data using template "
         "literals (the backtick-string syntax):")
add_code_block('''return `
  <div class="ro-page">
    <div class="ro-header">
      <div class="ro-company">${escHtml(shopInfo.name)}</div>
      ...
    </div>
    ...
  </div>
`;''')

heading("Two things to know", level=3)
add_para("Template literals — backticks let you embed ${expression} directly inside a string. "
         "${escHtml(shopInfo.name)} evaluates the function call and pastes the result into "
         "the string. Way cleaner than string concatenation with +.")

add_para("Why escHtml? Because we’re building HTML from user/API data. If a company name "
         "contained <script>alert('xss')</script>, raw interpolation would inject it as "
         "executable HTML. escHtml runs the string through a DOM textNode-then-innerHTML trick "
         "that converts < to &lt; etc., neutralizing it. This is the XSS protection.")

add_para("The job rows are built in a loop using .forEach, accumulating into a jobRows string "
         "that’s then injected into the table.", space_before=4)

# ============================================================
# SECTION 10
# ============================================================
heading("10. CSS and Print-Specific Styling", level=1)

add_para("The CSS is a big template literal stored in SHARED_STYLES, then concatenated with "
         "print-only overrides in PRINT_CSS. The script injects it into the page with:")
add_code_block('''function ensureStyles() {
  if (document.getElementById('wo-print-styles')) return;
  const style = document.createElement('style');
  style.id = 'wo-print-styles';
  style.textContent = PRINT_CSS;
  document.head.appendChild(style);
}''')

heading("The key print-specific tricks", level=3)
add_code_block('''@media print {
  body > *:not(#wo-print-overlay) { display: none !important; }
  #wo-print-overlay { display: block !important; }
  @page { size: letter; margin: 0.5in; }
}''', caption="@media print rules apply only when the browser is in print/PDF mode:")

add_para("What this does:")
add_bullet("Hide every direct child of <body> except our overlay")
add_bullet("Show our overlay")
add_bullet("Set the printed page to US Letter with 0.5-inch margins")

add_para("So when the user clicks Print, the entire MyGeotab UI disappears from the printout "
         "and only our Work Order shows. Once they close the print dialog, normal CSS takes "
         "over and MyGeotab reappears (we then remove the overlay element entirely as cleanup).",
         space_before=4)

heading("Page-break controls — for the parts table", level=3)
add_code_block('''.ro-section-jobs { page-break-inside: auto; }
.ro-table thead { display: table-header-group; }
.ro-table tbody tr { page-break-inside: avoid; }''')

add_para("Without display: table-header-group, only page 1 would show the column headers. With "
         "it, every page that contains table rows also gets the header row at the top. "
         "Critical for multi-page work orders.")

# ============================================================
# SECTION 11
# ============================================================
heading("11. Triggering the Print", level=1)

add_code_block('''function triggerPrint(html, previewMode, title) {
  ensureStyles();
  const overlay = document.createElement('div');
  overlay.id = 'wo-print-overlay';
  overlay.innerHTML = html;
  document.body.appendChild(overlay);

  document.title = title;        // becomes the default PDF filename
  window.print();                // opens the browser's print dialog

  window.addEventListener('afterprint', cleanup, { once: true });
  setTimeout(cleanup, 60000);    // fallback in case afterprint never fires
}''')

add_para("window.print() is a built-in browser function that opens the native print dialog. "
         "The browser already knows how to render HTML to PDF — it does this every time you "
         "“Save as PDF” from any web page. We just leverage that.")

add_para("document.title controls the filename when the user picks “Save as PDF” — that’s why "
         "renaming Repair Order to Work Order changed the default filename to Work Order — "
         "000019.pdf.")

add_para("afterprint fires when the user closes the print dialog (either after printing or "
         "cancelling). The cleanup removes our overlay element from the DOM and restores the "
         "original page title. The 60-second setTimeout is a safety net — in some browsers, "
         "afterprint is unreliable.")

# ============================================================
# SECTION 12
# ============================================================
heading("12. Keeping the Button Styled (MutationObserver)", level=1)

add_para("The MyGeotab add-in system creates a button for you automatically. But the styling "
         "is plain text — no icon, wrong color. To make it look like a proper primary action "
         "button (blue with a printer icon), we manually restyle it after it appears:")

add_code_block('''function applyButtonStyle() {
  var btn = document.querySelector('button[aria-label="Print"]');
  if (!btn) return false;
  var hasIcon = !!btn.querySelector('svg');
  if (btn.dataset.woStyled === '1' && hasIcon) return false;
  btn.classList.add('zen-button--primary', 'zen-caption', 'zen-text-icon-button');
  btn.innerHTML = PRINTER_SVG_ICON + '<span class="zen-caption__content">Print</span>';
  btn.dataset.woStyled = '1';
  return true;
}''')

add_bullet("We use MyGeotab’s own CSS classes (zen-button--primary, etc.) so the button "
           "matches the rest of the UI.")
add_bullet("We use a data-woStyled attribute as a marker so we don’t re-style the same button "
           "on every check.")

heading("The persistence problem", level=3)
add_para("MyGeotab is a Single-Page Application (SPA). When you navigate between work orders, "
         "MyGeotab re-renders the toolbar — sometimes replacing the button entirely, sometimes "
         "wiping its inner HTML while keeping the same element. Our marker survives on the "
         "same element, but a new element starts unmarked, and a content-wiped element loses "
         "our SVG icon.")

add_para("The fix is a MutationObserver:")
add_code_block('''function watchForButton() {
  applyButtonStyle();
  var queued = false;
  var observer = new MutationObserver(function () {
    if (queued) return;
    queued = true;
    requestAnimationFrame(function () {
      queued = false;
      applyButtonStyle();
    });
  });
  observer.observe(document.body, { childList: true, subtree: true });
}''')

add_para("A MutationObserver is a browser API that fires a callback whenever the DOM changes "
         "anywhere in the page. We listen on document.body with subtree: true to catch changes "
         "deep in the tree.")

add_callout("Why requestAnimationFrame?",
            "MutationObserver on the whole body can fire thousands of times per second on a "
            "busy SPA. To avoid spending CPU on redundant work, we coalesce with "
            "requestAnimationFrame — we only do one styling check per frame (roughly every "
            "16ms at 60fps), no matter how many mutations fire in between. The queued flag is "
            "a simple debounce.")

# ============================================================
# SECTION 13
# ============================================================
heading("13. Shop Info: localStorage + API + Prompt Fallback", level=1)

add_para("The header of the printed doc shows the shop’s name, address, and phone. We need "
         "that data, but it’s not on the work order itself. The script tries three sources in "
         "order:")
add_code_block('''async function resolveShopInfo(api) {
  // 1. localStorage override (user edited it before)
  const override = getShopInfoOverride();
  if (override && override.name) return override;

  // 2. API: CompanyDetails entity
  const apiInfo = await fetchShopInfoFromApi(api);
  if (apiInfo.name) {
    saveShopInfoOverride(apiInfo);  // cache for next time
    return apiInfo;
  }

  // 3. Prompt user to enter it manually
  return await promptShopInfo();
}''')

add_bullet_mixed([("localStorage", {'bold': True}),
                  (" is browser-built-in key/value storage. Survives across page loads but is "
                   "per-browser (so logging in on a different device starts fresh). We use it "
                   "as a cache.", {})])
add_bullet_mixed([("The CompanyDetails API call", {'bold': True}),
                  (" returns companyName and phoneNumber. There’s no address field exposed "
                   "there, which is why the prompt UI lets the user fill it in manually.", {})])
add_bullet_mixed([("The prompt", {'bold': True}),
                  (" is a custom modal overlay — pure JavaScript creating a div with inline "
                   "styles, appending to body, returning a Promise that resolves when the user "
                   "clicks Save or Cancel. No external library.", {})])

# ============================================================
# SECTION 14
# ============================================================
heading("14. The End-to-End Flow", level=1)

add_para("When the user clicks the Print button, this is everything that happens in order:")

flow = [
    "MyGeotab calls window.geotab.customButtons.printWorkOrder(event, nativeApi, state)",
    "We wrap nativeApi in a Promise-based interface",
    "We extract the work order ID from state.entity.id (or fall back to parsing the URL hash)",
    "We resolve shop info (localStorage → API → prompt)",
    "We multiCall the API for the WO + jobs in one network round trip",
    "We call the API for the device",
    "We pass all the data to buildRepairOrderHTML() which builds an HTML string",
    "We inject our CSS into the page (idempotent — only happens once)",
    "We append an overlay div containing the HTML to body",
    "We set document.title so the saved PDF gets a nice filename",
    "We call window.print() — browser opens its native print dialog",
    "Browser uses our @media print CSS to hide MyGeotab and show only our overlay",
    "User picks “Save as PDF” or a physical printer",
    "When they close the dialog, afterprint fires",
    "We remove the overlay element and restore the original title",
    "Page returns to normal MyGeotab state",
]
for i, step in enumerate(flow, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    p.paragraph_format.space_after = Pt(3)
    num = p.add_run(f"{i:>2}.  ")
    num.font.name = 'Consolas'
    num.font.size = Pt(10.5)
    num.font.bold = True
    num.font.color.rgb = BRAND_BLUE_DARK
    body = p.add_run(step)
    body.font.name = 'Calibri'
    body.font.size = Pt(11)
    body.font.color.rgb = INK

# ============================================================
# SECTION 15
# ============================================================
heading("15. The Dev Harness", level=1)

add_para("Tucked at the bottom:")
add_code_block('''window._woPrint = {
  resolveShopInfo,
  buildRepairOrderHTML,
  triggerPrint,
  fetchWorkOrderData,
  resetShopInfo: function () {
    localStorage.removeItem(STORAGE_KEY);
  }
};''')

add_para("This exposes the internal functions to the browser console under window._woPrint. "
         "So if you want to test something without clicking the button, you can open DevTools "
         "and run:")
add_code_block('''_woPrint.resetShopInfo();      // clear cached shop info, force re-fetch
_woPrint.triggerPrint('<div>test</div>', true);  // preview overlay without printing''')

add_para("It’s purely a development aid. Production users never touch it.",
         italic=True, color=GREY_MID, space_before=4)

# ============================================================
# SECTION 16
# ============================================================
heading("16. Tech Stack Summary", level=1)

add_three_col_table([
    ("Runtime", "Vanilla JavaScript (ES2017+)",
     "No bundler, no framework — must run as-is in the browser. MyGeotab loads it like a <script> tag."),
    ("Async", "async/await over Promises",
     "Wraps MyGeotab’s callback API for clean control flow."),
    ("DOM", "createElement, innerHTML, appendChild",
     "No React, no Vue. Plain DOM manipulation. The doc is throwaway HTML, no reactivity needed."),
    ("Styling", "Inline CSS injected via <style> tag",
     "Self-contained — no separate CSS file to ship. @media print for print-only rules."),
    ("Persistence", "localStorage",
     "For caching shop info between sessions."),
    ("Network", "MyGeotab’s pre-authed nativeApi",
     "No fetch, no auth headers, no CORS. The platform handles all of it."),
    ("PDF generation", "window.print() (browser-native)",
     "The browser already knows how to render HTML → PDF. Zero dependencies."),
    ("Lifecycle", "MutationObserver + requestAnimationFrame",
     "Watches MyGeotab’s SPA re-renders to keep our button styled."),
    ("Platform", "window.geotab.customButtons.<name>",
     "MyGeotab’s official add-in callback convention."),
], widths=(Inches(1.2), Inches(2.0), Inches(3.3)),
   header=("Layer", "What we use", "Why"))

add_callout("Zero external dependencies",
            "No npm, no webpack, no jQuery, no React. Just one .js file (~30KB) and one "
            "config.json (~340 bytes) uploaded to MyGeotab. That’s the entire deployable.",
            color_hex="EAF7EA",
            label_color=RGBColor(0x1F, 0x7A, 0x1F))

# ============================================================
# SAVE
# ============================================================
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"Wrote: {OUT}")
print(f"Size: {OUT.stat().st_size:,} bytes")
